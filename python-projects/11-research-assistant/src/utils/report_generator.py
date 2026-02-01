"""
Report Generator for Research Assistant.

Generates research reports in multiple formats:
- Markdown (.md)
- HTML (.html)
- JSON (.json)
- PDF (.pdf) - optional, requires weasyprint
- DOCX (.docx) - optional, requires python-docx
"""

import logging
import json
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path


class ReportGenerator:
    """Generates research reports in multiple formats."""

    def __init__(self, output_dir: str = './data/output'):
        """
        Initialize report generator.

        Args:
            output_dir: Directory for generated reports
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        logging.info(f"ReportGenerator initialized (output: {self.output_dir})")

    def generate_report(
        self,
        research_data: Dict[str, Any],
        format: str = 'markdown',
        filename: Optional[str] = None
    ) -> Path:
        """
        Generate research report in specified format.

        Args:
            research_data: Research results dictionary
            format: Output format ('markdown', 'html', 'json', 'pdf', 'docx')
            filename: Optional custom filename (without extension)

        Returns:
            Path to generated report file
        """
        format = format.lower()

        # Generate filename if not provided
        if not filename:
            query_short = research_data['query'][:50].replace(' ', '_')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"research_{query_short}_{timestamp}"

        # Clean filename
        filename = self._sanitize_filename(filename)

        # Generate report based on format
        if format == 'markdown' or format == 'md':
            return self._generate_markdown(research_data, filename)
        elif format == 'html':
            return self._generate_html(research_data, filename)
        elif format == 'json':
            return self._generate_json(research_data, filename)
        elif format == 'pdf':
            return self._generate_pdf(research_data, filename)
        elif format == 'docx' or format == 'doc':
            return self._generate_docx(research_data, filename)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_markdown(
        self,
        data: Dict[str, Any],
        filename: str
    ) -> Path:
        """Generate Markdown report."""
        output_path = self.output_dir / f"{filename}.md"

        content = []

        # Title
        content.append(f"# Research Report: {data['query']}\n")

        # Metadata
        content.append("## Metadata\n")
        content.append(f"- **Query ID**: `{data.get('query_id', 'N/A')}`")
        content.append(f"- **Generated**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        if 'stats' in data:
            stats = data['stats']
            content.append(f"- **Sources Used**: {stats.get('used_sources', 0)}")
            content.append(f"- **Findings**: {stats.get('findings', 0)}")
            content.append(f"- **Average Confidence**: {stats.get('avg_confidence', 0):.2f}")
            content.append(f"- **Processing Time**: {stats.get('processing_time', 0):.1f}s")

        content.append("\n---\n")

        # Summary
        content.append("## Summary\n")
        content.append(data.get('summary', 'No summary available.'))
        content.append("\n")

        # Findings
        if data.get('findings'):
            content.append("## Key Findings\n")
            for i, finding in enumerate(data['findings'], 1):
                content.append(f"### {i}. {finding.get('type', 'Finding').title()}\n")
                content.append(finding.get('text', ''))
                content.append(f"\n**Confidence**: {finding.get('confidence', 0):.2f} | ")
                content.append(f"**Sources**: {finding.get('sources', 0)}\n")

        # Sources
        if data.get('sources'):
            content.append("## Sources\n")
            for i, source in enumerate(data['sources'], 1):
                content.append(f"{i}. **{source.get('title', 'Untitled')}**")
                if source.get('url'):
                    content.append(f" - [{source['url']}]({source['url']})")
                content.append(f" (*{source.get('type', 'unknown')}*)")
                if 'composite_score' in source:
                    content.append(f" - Score: {source['composite_score']:.2f}")
                content.append("")

        # Citations
        if data.get('citations'):
            content.append("\n## References\n")
            for i, citation in enumerate(data['citations'], 1):
                content.append(f"{i}. {citation}")
                content.append("")

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))

        logging.info(f"Generated Markdown report: {output_path}")
        return output_path

    def _generate_html(
        self,
        data: Dict[str, Any],
        filename: str
    ) -> Path:
        """Generate HTML report."""
        output_path = self.output_dir / f"{filename}.html"

        html = []

        # HTML header
        html.append("<!DOCTYPE html>")
        html.append("<html lang='en'>")
        html.append("<head>")
        html.append("  <meta charset='UTF-8'>")
        html.append(f"  <title>Research Report: {self._escape_html(data['query'])}</title>")
        html.append("  <style>")
        html.append(self._get_html_styles())
        html.append("  </style>")
        html.append("</head>")
        html.append("<body>")

        # Header
        html.append(f"  <h1>Research Report: {self._escape_html(data['query'])}</h1>")

        # Metadata
        html.append("  <div class='metadata'>")
        html.append(f"    <p><strong>Query ID:</strong> {data.get('query_id', 'N/A')}</p>")
        html.append(f"    <p><strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")

        if 'stats' in data:
            stats = data['stats']
            html.append(f"    <p><strong>Sources Used:</strong> {stats.get('used_sources', 0)}</p>")
            html.append(f"    <p><strong>Findings:</strong> {stats.get('findings', 0)}</p>")
            html.append(f"    <p><strong>Average Confidence:</strong> {stats.get('avg_confidence', 0):.2f}</p>")
            html.append(f"    <p><strong>Processing Time:</strong> {stats.get('processing_time', 0):.1f}s</p>")

        html.append("  </div>")

        # Summary
        html.append("  <h2>Summary</h2>")
        html.append(f"  <div class='summary'>{self._escape_html(data.get('summary', 'No summary available.'))}</div>")

        # Findings
        if data.get('findings'):
            html.append("  <h2>Key Findings</h2>")
            for i, finding in enumerate(data['findings'], 1):
                html.append("  <div class='finding'>")
                html.append(f"    <h3>{i}. {self._escape_html(finding.get('type', 'Finding').title())}</h3>")
                html.append(f"    <p>{self._escape_html(finding.get('text', ''))}</p>")
                html.append(f"    <p class='finding-meta'>")
                html.append(f"      <span class='confidence'>Confidence: {finding.get('confidence', 0):.2f}</span> | ")
                html.append(f"      <span class='sources'>Sources: {finding.get('sources', 0)}</span>")
                html.append(f"    </p>")
                html.append("  </div>")

        # Sources
        if data.get('sources'):
            html.append("  <h2>Sources</h2>")
            html.append("  <ol class='sources'>")
            for source in data['sources']:
                html.append("    <li>")
                html.append(f"      <strong>{self._escape_html(source.get('title', 'Untitled'))}</strong>")
                if source.get('url'):
                    html.append(f" - <a href='{source['url']}' target='_blank'>{source['url']}</a>")
                html.append(f" <em>({source.get('type', 'unknown')})</em>")
                if 'composite_score' in source:
                    html.append(f" - Score: {source['composite_score']:.2f}")
                html.append("    </li>")
            html.append("  </ol>")

        # Citations
        if data.get('citations'):
            html.append("  <h2>References</h2>")
            html.append("  <ol class='citations'>")
            for citation in data['citations']:
                html.append(f"    <li>{self._escape_html(citation)}</li>")
            html.append("  </ol>")

        # Footer
        html.append("</body>")
        html.append("</html>")

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html))

        logging.info(f"Generated HTML report: {output_path}")
        return output_path

    def _generate_json(
        self,
        data: Dict[str, Any],
        filename: str
    ) -> Path:
        """Generate JSON report."""
        output_path = self.output_dir / f"{filename}.json"

        # Add generation metadata
        report_data = {
            'generated_at': datetime.now().isoformat(),
            'query': data.get('query'),
            'query_id': data.get('query_id'),
            'summary': data.get('summary'),
            'findings': data.get('findings', []),
            'sources': data.get('sources', []),
            'citations': data.get('citations', []),
            'stats': data.get('stats', {})
        }

        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False)

        logging.info(f"Generated JSON report: {output_path}")
        return output_path

    def _generate_pdf(
        self,
        data: Dict[str, Any],
        filename: str
    ) -> Path:
        """Generate PDF report (requires weasyprint)."""
        try:
            from weasyprint import HTML
        except ImportError:
            logging.error("PDF generation requires weasyprint: pip install weasyprint")
            raise ValueError("PDF generation not available (install weasyprint)")

        # Generate HTML first
        html_path = self._generate_html(data, f"{filename}_temp")

        # Convert HTML to PDF
        output_path = self.output_dir / f"{filename}.pdf"

        HTML(filename=str(html_path)).write_pdf(output_path)

        # Clean up temp HTML
        html_path.unlink()

        logging.info(f"Generated PDF report: {output_path}")
        return output_path

    def _generate_docx(
        self,
        data: Dict[str, Any],
        filename: str
    ) -> Path:
        """Generate DOCX report (requires python-docx)."""
        try:
            from docx import Document
            from docx.shared import RGBColor
        except ImportError:
            logging.error("DOCX generation requires python-docx: pip install python-docx")
            raise ValueError("DOCX generation not available (install python-docx)")

        output_path = self.output_dir / f"{filename}.docx"

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(f"Research Report: {data['query']}", level=1)
        title_format = title.runs[0].font
        title_format.color.rgb = RGBColor(44, 62, 80)

        # Metadata
        doc.add_heading('Metadata', level=2)
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Query ID: ").bold = True
        metadata_para.add_run(f"{data.get('query_id', 'N/A')}\n")
        metadata_para.add_run("Generated: ").bold = True
        metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        if 'stats' in data:
            stats = data['stats']
            metadata_para.add_run("Sources Used: ").bold = True
            metadata_para.add_run(f"{stats.get('used_sources', 0)}\n")
            metadata_para.add_run("Findings: ").bold = True
            metadata_para.add_run(f"{stats.get('findings', 0)}\n")
            metadata_para.add_run("Average Confidence: ").bold = True
            metadata_para.add_run(f"{stats.get('avg_confidence', 0):.2f}\n")
            metadata_para.add_run("Processing Time: ").bold = True
            metadata_para.add_run(f"{stats.get('processing_time', 0):.1f}s\n")

        # Summary
        doc.add_heading('Summary', level=2)
        summary_para = doc.add_paragraph(data.get('summary', 'No summary available.'))
        summary_para.style = 'Intense Quote'

        # Findings
        if data.get('findings'):
            doc.add_heading('Key Findings', level=2)
            for i, finding in enumerate(data['findings'], 1):
                doc.add_heading(f"{i}. {finding.get('type', 'Finding').title()}", level=3)
                doc.add_paragraph(finding.get('text', ''))

                meta_para = doc.add_paragraph()
                meta_para.add_run("Confidence: ").bold = True
                confidence_run = meta_para.add_run(f"{finding.get('confidence', 0):.2f}")
                confidence_run.font.color.rgb = RGBColor(46, 204, 113)
                meta_para.add_run(" | ")
                meta_para.add_run("Sources: ").bold = True
                meta_para.add_run(f"{finding.get('sources', 0)}")

        # Sources
        if data.get('sources'):
            doc.add_heading('Sources', level=2)
            for i, source in enumerate(data['sources'], 1):
                source_para = doc.add_paragraph(style='List Number')
                source_para.add_run(source.get('title', 'Untitled')).bold = True

                if source.get('url'):
                    source_para.add_run(f" - ")
                    url_run = source_para.add_run(source['url'])
                    url_run.font.color.rgb = RGBColor(52, 152, 219)

                source_para.add_run(f" ({source.get('type', 'unknown')})")

                if 'composite_score' in source:
                    source_para.add_run(f" - Score: {source['composite_score']:.2f}")

        # Citations
        if data.get('citations'):
            doc.add_heading('References', level=2)
            for i, citation in enumerate(data['citations'], 1):
                doc.add_paragraph(citation, style='List Number')

        # Save document
        doc.save(str(output_path))

        logging.info(f"Generated DOCX report: {output_path}")
        return output_path

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for filesystem."""
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')

        # Limit length
        if len(filename) > 100:
            filename = filename[:100]

        return filename

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        if not text:
            return ''

        text = str(text)
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&#39;')
        return text

    def _get_html_styles(self) -> str:
        """Get CSS styles for HTML report."""
        return """
    body {
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, Cantarell, sans-serif;
      line-height: 1.6;
      max-width: 900px;
      margin: 0 auto;
      padding: 20px;
      color: #333;
    }

    h1 {
      color: #2c3e50;
      border-bottom: 3px solid #3498db;
      padding-bottom: 10px;
    }

    h2 {
      color: #34495e;
      border-bottom: 2px solid #ecf0f1;
      padding-bottom: 5px;
      margin-top: 30px;
    }

    h3 {
      color: #7f8c8d;
    }

    .metadata {
      background: #ecf0f1;
      padding: 15px;
      border-radius: 5px;
      margin: 20px 0;
    }

    .metadata p {
      margin: 5px 0;
    }

    .summary {
      background: #fff;
      padding: 20px;
      border-left: 4px solid #3498db;
      margin: 20px 0;
    }

    .finding {
      background: #f9f9f9;
      padding: 15px;
      margin: 15px 0;
      border-radius: 5px;
      border-left: 4px solid #2ecc71;
    }

    .finding-meta {
      color: #7f8c8d;
      font-size: 0.9em;
      margin-top: 10px;
    }

    .confidence {
      font-weight: bold;
      color: #2ecc71;
    }

    .sources {
      list-style-type: decimal;
      padding-left: 20px;
    }

    .sources li {
      margin: 10px 0;
    }

    .citations {
      list-style-type: decimal;
      padding-left: 20px;
    }

    .citations li {
      margin: 10px 0;
      color: #555;
    }

    a {
      color: #3498db;
      text-decoration: none;
    }

    a:hover {
      text-decoration: underline;
    }
"""

    def get_info(self) -> Dict[str, Any]:
        """Get information about the report generator."""
        # Check if PDF generation is available
        pdf_available = False
        try:
            import weasyprint
            pdf_available = True
        except ImportError:
            pass

        # Check if DOCX generation is available
        docx_available = False
        try:
            import docx
            docx_available = True
        except ImportError:
            pass

        # Build supported formats list
        formats = ['markdown', 'html', 'json']
        if pdf_available:
            formats.append('pdf')
        if docx_available:
            formats.append('docx')

        return {
            'output_dir': str(self.output_dir),
            'supported_formats': formats,
            'pdf_available': pdf_available,
            'docx_available': docx_available
        }
