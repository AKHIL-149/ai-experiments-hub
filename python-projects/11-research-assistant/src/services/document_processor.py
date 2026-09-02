"""
Document Processor for Research Assistant.

Extracts text from uploaded documents (PDF, TXT, MD, DOCX) and splits
it into overlapping chunks suitable for embedding + retrieval.
"""

import logging
from pathlib import Path
from typing import List

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("PyMuPDF not installed. PDF extraction will not work.")

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX extraction will not work.")


SUPPORTED_TYPES = {'pdf', 'txt', 'md', 'docx'}


class DocumentProcessingError(Exception):
    """Raised when a document can't be read or extracted."""


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract plain text from a document file.

    Args:
        file_path: Path to the file on disk
        file_type: One of 'pdf', 'txt', 'md', 'docx' (case-insensitive)

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If the type is unsupported or extraction fails
    """
    file_type = file_type.lower().lstrip('.')

    if file_type not in SUPPORTED_TYPES:
        raise DocumentProcessingError(f"Unsupported file type: {file_type}")

    try:
        if file_type == 'pdf':
            return _extract_pdf(file_path)
        elif file_type in ('txt', 'md'):
            return _extract_plain_text(file_path)
        elif file_type == 'docx':
            return _extract_docx(file_path)
    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(f"Failed to extract text: {e}") from e


def _extract_pdf(file_path: str) -> str:
    if not PYMUPDF_AVAILABLE:
        raise DocumentProcessingError("PDF extraction requires PyMuPDF (fitz)")

    doc = fitz.open(file_path)
    try:
        text_parts = [page.get_text() for page in doc]
    finally:
        doc.close()

    text = '\n\n'.join(t for t in text_parts if t.strip())
    if not text.strip():
        raise DocumentProcessingError("No extractable text found in PDF (may be scanned/image-only)")
    return text


def _extract_plain_text(file_path: str) -> str:
    # Try common encodings rather than assuming UTF-8 - user-uploaded
    # text files are not guaranteed to be UTF-8.
    for encoding in ('utf-8', 'utf-8-sig', 'latin-1'):
        try:
            return Path(file_path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError("Could not decode text file with any supported encoding")


def _extract_docx(file_path: str) -> str:
    if not DOCX_AVAILABLE:
        raise DocumentProcessingError("DOCX extraction requires python-docx")

    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables - a document that's mostly a table
    # (common for structured reports) would otherwise extract as empty.
    for table in document.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    text = '\n\n'.join(paragraphs)
    if not text.strip():
        raise DocumentProcessingError("No extractable text found in DOCX")
    return text


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 150
) -> List[str]:
    """
    Split text into overlapping chunks for embedding.

    Splits on paragraph boundaries where possible so chunks don't cut
    sentences in half arbitrarily; falls back to a hard character split
    for any single paragraph longer than chunk_size on its own.

    Args:
        text: Full extracted document text
        chunk_size: Target maximum characters per chunk
        overlap: Characters of overlap carried into the next chunk, so a
            fact split across a chunk boundary isn't lost entirely from
            either chunk's context

    Returns:
        List of text chunks (never empty strings)
    """
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    if not paragraphs:
        return []

    chunks = []
    current = ''

    for para in paragraphs:
        # A single paragraph longer than chunk_size on its own: hard-split it.
        if len(para) > chunk_size:
            if current:
                chunks.append(current)
                current = ''
            for i in range(0, len(para), chunk_size - overlap):
                chunks.append(para[i:i + chunk_size])
            continue

        candidate = f"{current}\n\n{para}" if current else para
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)
            # Carry the tail of the previous chunk forward as overlap context.
            tail = current[-overlap:] if current and overlap > 0 else ''
            current = f"{tail}\n\n{para}" if tail else para

    if current:
        chunks.append(current)

    return chunks
